"""DOCX renderer — A4 professional report format (python-docx).

Public API
----------
render_docx(report, output_path, template_id="bmj") -> Path
DocxRenderer().render(report, output_path, template="bmj") -> Path
"""
from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from ..templates.base import BaseTemplate


class DocxRenderer:
    """Thin class wrapper for CLI compatibility."""
    def render(self, report, output_path, template: str = "bmj") -> Path:
        return render_docx(report, output_path, template_id=template)


def render_docx(report, output_path: str | Path, template_id: str = "bmj") -> Path:
    """Render CritiqueReport as a professional A4 .docx file. Returns path."""
    try:
        from docx import Document
        from docx.shared import Cm, Pt
    except ImportError as exc:
        raise ImportError("python-docx required: pip install python-docx") from exc

    from .layout import HEX

    tmpl = BaseTemplate.all_templates().get(template_id)
    if tmpl is None:
        raise ValueError(f"Unknown template: {template_id!r}")

    sections_data = tmpl.build(report)
    doc  = Document()
    path = Path(output_path)

    # A4 page + margins
    doc_sec = doc.sections[0]
    doc_sec.page_width   = Cm(21.0)
    doc_sec.page_height  = Cm(29.7)
    doc_sec.left_margin  = Cm(2.5)
    doc_sec.right_margin = Cm(2.5)
    doc_sec.top_margin   = Cm(2.0)
    doc_sec.bottom_margin = Cm(2.5)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # color shortcuts
    navy = _rgb(HEX["primary"])
    mnav = _rgb(HEX["secondary"])

    # COVER PAGE ============================================================
    omega_str = (
        f"{report.omega_score:.4f}"
        + (f"  (hybrid {report.hybrid_omega:.4f})" if report.hybrid_omega is not None else "")
    )
    cover_meta: list[tuple[str, str]] = [
        ("PRECHECK MODE",    report.precheck.mode.value),
        ("CRITIQUE ROUND",   str(report.round_number)),
        ("OMEGA SCORE",      omega_str),
        ("NOT TRACEABLE",    f"{report.not_traceable_count()} finding(s)"),
        ("PARTIALLY TRACE.", f"{report.partially_traceable_count()} finding(s)"),
        ("GENERATED",        date.today().isoformat()),
    ]
    if report.experiment_class:
        cover_meta.insert(2, ("EXPERIMENT CLASS", report.experiment_class.value))

    _add_cover(doc, tmpl.DISPLAY_NAME, cover_meta, navy, mnav)
    doc.add_page_break()

    # BODY SECTIONS =========================================================
    _SKIP = {"Cover Page", "Title Page"}
    for sec in sections_data:
        if sec.title in _SKIP:
            continue
        _add_section(doc, sec, mnav)

    # SCORE APPENDIX ========================================================
    if report.irf_scores:
        sc = report.irf_scores
        dims = [
            ("M",         f"{sc.M:.3f}",         "Methodic Doubt"),
            ("A",         f"{sc.A:.3f}",         "Axiom / Hypothesis"),
            ("D",         f"{sc.D:.3f}",         "Deduction"),
            ("I",         f"{sc.I:.3f}",         "Induction"),
            ("F",         f"{sc.F:.3f}",         "Falsification"),
            ("P",         f"{sc.P:.3f}",         "Paradigm"),
            ("COMPOSITE", f"{sc.composite:.3f}", "PASS" if sc.passed else "WARN"),
        ]
        _add_dim_table(doc, "IRF-Calc 6D Score (LOGOS)", dims, mnav)

    if report.hsta_scores:
        h = report.hsta_scores
        dims = [
            ("N",         f"{h.N:.3f}",         "Novelty — semantic uniqueness"),
            ("C",         f"{h.C:.3f}",         "Consistency — cross-reference agreement"),
            ("T",         f"{h.T:.3f}",         "Temporality — citation recency"),
            ("R",         f"{h.R:.3f}",         "Reproducibility — method completeness"),
            ("COMPOSITE", f"{h.composite:.3f}", "Arithmetic mean (N+C+T+R)/4"),
        ]
        _add_dim_table(doc, "HSTA 4D Score (BioMedical-Paper-Harvester)", dims, mnav)

    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def _add_cover(doc, display_name: str, meta_rows, navy, mnav) -> None:
    from docx.shared import Pt

    p = doc.add_paragraph()
    r = p.add_run("VERITAS — AI Critique Experimental Report Analysis Framework")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = navy
    p.paragraph_format.space_after = Pt(4)

    sub = doc.add_paragraph()
    sub.add_run(display_name).font.color.rgb = mnav
    sub.runs[0].font.size = Pt(12)
    sub.paragraph_format.space_after = Pt(16)

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(meta_rows):
        row = table.add_row()
        kr = row.cells[0].paragraphs[0].add_run(k)
        kr.bold = True
        kr.font.size = Pt(9)
        kr.font.color.rgb = _rgb_white()
        _shade(row.cells[0], "16213e" if i % 2 == 0 else "0f3460")
        vr = row.cells[1].paragraphs[0].add_run(v)
        vr.font.size = Pt(9)
        _shade(row.cells[1], "f0f4f8" if i % 2 == 0 else "ffffff")
    doc.add_paragraph()


def _add_section(doc, sec, hdr_color) -> None:
    from docx.shared import Pt

    h = doc.add_heading(sec.title, level=1)
    if h.runs:
        h.runs[0].font.color.rgb = hdr_color
        h.runs[0].font.size = Pt(11)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after  = Pt(4)

    for para_text in sec.body.split("\n\n"):
        t = para_text.strip()
        if t:
            p = doc.add_paragraph()
            _bold_runs(p, t)
            p.paragraph_format.space_after = Pt(4)

    if sec.findings:
        _add_findings_table(doc, sec.findings)
    doc.add_paragraph()


def _add_findings_table(doc, findings: list[str]) -> None:
    from docx.shared import Pt

    _TC_COLORS = {
        "TRACEABLE":           "2d6a4f",
        "PARTIALLY TRACEABLE": "9c6644",
        "NOT TRACEABLE":       "c1121f",
    }
    pat = re.compile(r'^\[([^\]]+)\]\s+\[([^\]]+)\]\s+(.+)$')

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for cell, label in zip(table.rows[0].cells, ("CODE", "TRACEABILITY", "DESCRIPTION"), strict=False):
        r = cell.paragraphs[0].add_run(label)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = _rgb_white()
        _shade(cell, "0f3460")

    for i, raw in enumerate(findings):
        m = pat.match(raw.strip())
        code, tc, desc = (m.group(1), m.group(2), m.group(3)) if m else ("", "", raw)
        row = table.add_row()
        _fill(row.cells[0], code, Pt(8.5))
        _fill(row.cells[1], tc,   Pt(8.5),
              _rgb(f"#{_TC_COLORS.get(tc.upper(), '212529')}"), bold=True)
        _fill(row.cells[2], desc, Pt(8.5))
        if i % 2 == 0:
            for cell in row.cells:
                _shade(cell, "f0f4f8")


def _add_dim_table(doc, title: str, dims: list[tuple[str, str, str]], hdr_color) -> None:
    from docx.shared import Pt

    h = doc.add_heading(title, level=1)
    if h.runs:
        h.runs[0].font.color.rgb = hdr_color
        h.runs[0].font.size = Pt(11)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for cell, label in zip(table.rows[0].cells, ("DIM", "SCORE", "MEANING"), strict=False):
        r = cell.paragraphs[0].add_run(label)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = _rgb_white()
        _shade(cell, "16213e")

    for i, (d, v, m) in enumerate(dims):
        is_composite = d == "COMPOSITE"
        row = table.add_row()
        _fill(row.cells[0], d, Pt(8.5), bold=is_composite)
        _fill(row.cells[1], v, Pt(8.5), bold=is_composite)
        _fill(row.cells[2], m, Pt(8.5), bold=is_composite)
        if not is_composite and i % 2 == 0:
            for cell in row.cells:
                _shade(cell, "f0f4f8")
        if is_composite:
            for cell in row.cells:
                _shade(cell, "0f3460")
                cell.paragraphs[0].runs[0].font.color.rgb = _rgb_white()
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Micro-helpers (cell shading, text runs)
# ---------------------------------------------------------------------------

def _rgb(hex_str: str):
    from docx.shared import RGBColor
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _rgb_white():
    from docx.shared import RGBColor
    return RGBColor(0xFF, 0xFF, 0xFF)


def _shade(cell, hex_bg: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_bg.lstrip("#"))
    tc_pr.append(shd)


def _fill(cell, text: str, size, color=None, bold: bool = False) -> None:
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.size = size
    r.bold = bold
    if color:
        r.font.color.rgb = color


def _bold_runs(para, text: str) -> None:
    """Add text to paragraph, rendering **bold** markers."""
    parts = re.split(r'\*\*(.+?)\*\*', text)
    for i, part in enumerate(parts):
        if part:
            para.add_run(part).bold = (i % 2 == 1)
