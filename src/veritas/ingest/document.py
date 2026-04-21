"""Document ingestion layer — wraps Flamehaven-Filesearch file_parser + text_chunker."""

from __future__ import annotations

import sys
from pathlib import Path

_FF_ROOT = r"D:\Sanctum\Flamehaven-Filesearch"

SUPPORTED = {".pdf", ".docx", ".doc", ".txt", ".md", ".rtf", ".html", ".tex"}


def extract_text(file_path: str | Path, use_cache: bool = False) -> str:
    """Extract plain text from a supported document.

    Delegates to Flamehaven-Filesearch's BackendRegistry when available,
    then falls back to pymupdf → python-docx → plain read.
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix not in SUPPORTED:
        raise ValueError(f"Unsupported file type: {suffix}. Supported: {SUPPORTED}")

    # Primary: Flamehaven-Filesearch
    text = _try_flamehaven(str(path), use_cache)
    if text:
        return text

    # Fallback chain
    if suffix == ".pdf":
        text = _try_pymupdf(path) or _try_pypdf(path)
    elif suffix in {".docx", ".doc"}:
        text = _try_docx(path)
    else:
        text = path.read_text(encoding="utf-8", errors="replace")

    return text or ""


def extract_chunks(
    file_path: str | Path,
    chunk_size: int = 512,
    overlap: int = 50,
) -> list[dict]:
    """Extract text and return structured chunks for RAG indexing.

    Each chunk: {"text": str, "chunk_index": int, "source": str}
    """
    text = extract_text(file_path)
    if not text.strip():
        return []

    chunks = _chunk(text, chunk_size, overlap)
    source = Path(file_path).name
    return [{"text": c, "chunk_index": i, "source": source} for i, c in enumerate(chunks)]


# ── private helpers ────────────────────────────────────────────────────────────


def _try_flamehaven(path: str, use_cache: bool) -> str | None:
    try:
        if _FF_ROOT not in sys.path:
            sys.path.insert(0, _FF_ROOT)
        from flamehaven_filesearch.engine.file_parser import extract_text as ff_extract

        return ff_extract(path, use_cache=use_cache) or None
    except Exception:
        return None


def _try_pymupdf(path: Path) -> str | None:
    try:
        import fitz

        doc = fitz.open(str(path))
        return "\n".join(page.get_text() for page in doc) or None
    except Exception:
        return None


def _try_pypdf(path: Path) -> str | None:
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        return "\n".join(p.extract_text() or "" for p in reader.pages) or None
    except Exception:
        return None


def _try_docx(path: Path) -> str | None:
    try:
        from docx import Document

        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts) or None
    except Exception:
        return None


def _chunk(text: str, size: int, overlap: int) -> list[str]:
    """Delegate to Flamehaven chunker when available, else simple word-based split."""
    try:
        if _FF_ROOT not in sys.path:
            sys.path.insert(0, _FF_ROOT)
        from flamehaven_filesearch.engine.text_chunker import chunk_text

        chunks = chunk_text(text, max_tokens=size, overlap_tokens=overlap)
        return [c["text"] for c in chunks] if chunks and isinstance(chunks[0], dict) else chunks
    except Exception:
        pass  # flamehaven_filesearch not installed; fall through to built-in word splitter

    words = text.split()
    result, i = [], 0
    while i < len(words):
        result.append(" ".join(words[i : i + size]))
        i += max(1, size - overlap)
    return result
